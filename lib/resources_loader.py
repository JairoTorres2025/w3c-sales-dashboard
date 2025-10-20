from __future__ import annotations

import re
from typing import List, Dict
from pathlib import Path

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None  # type: ignore


_HEADERS = {"tool", "description", "link", "how-to", "how to", "howto"}
_URL_RE = re.compile(r"https?://\S+", re.I)


def _norm_header(s: str) -> str:
    s = (s or "").strip().lower()
    s = s.replace(" ", "").replace("-", "")
    if s in {"howto", "howto"}:
        return "howto"
    return s


def _extract_url(text: str) -> str:
    if not text:
        return ""
    m = _URL_RE.search(text)
    return m.group(0) if m else ""


def parse_tools_docx(path: str) -> List[Dict[str, str]]:
    p = Path(path)
    if Document is None or not p.exists():
        return []
    doc = Document(str(p))

    # Prefer a table with expected headers
    for tbl in doc.tables:
        headers = [_norm_header(c.text) for c in tbl.rows[0].cells]
        if not headers:
            continue
        hdr_set = set(headers)
        if {"tool", "description"}.issubset(hdr_set):
            # map columns
            idx = {h: i for i, h in enumerate(headers)}
            def get(cell_name: str, row_cells):
                i = idx.get(cell_name)
                return row_cells[i].text.strip() if i is not None and i < len(row_cells) else ""
            entries: List[Dict[str, str]] = []
            for r in tbl.rows[1:]:
                cells = r.cells
                tool = get("tool", cells)
                if not tool:
                    continue
                desc = get("description", cells)
                howto = get("howto", cells) or get("how-to", cells) or get("how to", cells)
                link = get("link", cells)
                if not link:
                    # try to find url inside other fields
                    link = _extract_url(desc) or _extract_url(howto)
                entries.append({
                    "Tool": tool,
                    "Description": desc,
                    "Link": link,
                    "How-to": howto,
                })
            if entries:
                return entries

    # Fallback: build entries from paragraphs separated by blank lines
    paras = [p.text.strip() for p in doc.paragraphs]
    blocks: List[List[str]] = []
    cur: List[str] = []
    for t in paras:
        if not t:
            if cur:
                blocks.append(cur); cur = []
        else:
            cur.append(t)
    if cur:
        blocks.append(cur)

    entries: List[Dict[str, str]] = []
    for b in blocks:
        if not b:
            continue
        tool = b[0]
        text = "\n".join(b[1:])
        link = _extract_url(text)
        entries.append({
            "Tool": tool,
            "Description": text.split("\n")[0][:400],
            "Link": link,
            "How-to": text[:800],
        })
    return entries
